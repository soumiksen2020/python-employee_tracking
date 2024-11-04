import os
import time
import threading
from datetime import datetime, timedelta
import pymysql
import pyautogui
import pystray
from PIL import Image
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np

# Database configuration
DB_CONFIG = {
    'user': 'root',
    'password': '',
    'host': 'localhost',
    'database': 'usage_tracker'
}

# Global variables for time tracking
tracking = False
start_time = None
work_duration = timedelta(0)
screenshot_dir = "screenshots"
interval = 10  # Screenshot interval in seconds

# Step 1: Fetch application usage data from the database
def fetch_usage_data():
    try:
        conn = pymysql.connect(**DB_CONFIG)
        with conn.cursor() as cursor:
            cursor.execute("SELECT app_name, start_time, end_time, duration FROM applications")
            logs = cursor.fetchall()
            print("Data fetched from database:", logs)
    except pymysql.MySQLError as err:
        print(f"Error fetching data: {err}")
        return []
    finally:
        if conn:
            conn.close()
    return logs

# Step 2: Take a screenshot and save it as JPEG in the screenshots directory
def take_screenshot():
    screenshot_name = f"screenshot_{datetime.now().strftime('%Y%m%d_%H%M%S')}.jpeg"
    screenshot_path = os.path.join(screenshot_dir, screenshot_name)
    screenshot = pyautogui.screenshot()
    screenshot = screenshot.convert("RGB")
    screenshot.save(screenshot_path, "JPEG")
    print(f"Screenshot saved as: {screenshot_path}")

# Step 3: Generate a Word report with a usage chart
# Step 3: Generate a Word report with a usage chart, including login/logout times and active duration
# Step 3: Generate a Word report with a usage chart, including login/logout times and active duration
def generate_word_report(logs):
    if not logs:
        print("No logs found to generate the report.")
        return

    # Initialize a new Word document
    doc = Document()
    doc.add_heading("Application Usage Report", 0)

    # Add a table for the application logs
    doc.add_heading("Usage Details:", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'App Name'
    hdr_cells[1].text = 'Start Time'
    hdr_cells[2].text = 'End Time'
    hdr_cells[3].text = 'Duration (seconds)'

    # Populate the table with logs and accumulate durations for charting
    app_durations = {}
    for log in logs:
        # Exclude applications with zero duration
        if log[3] <= 0:
            continue

        row_cells = table.add_row().cells
        row_cells[0].text = log[0]
        row_cells[1].text = log[1].strftime("%Y-%m-%d %H:%M:%S")
        row_cells[2].text = log[2].strftime("%Y-%m-%d %H:%M:%S")
        row_cells[3].text = str(log[3])

        # Accumulate duration for each application
        app_durations[log[0]] = app_durations.get(log[0], 0) + log[3]

    # Step 4: Create a bar chart for application usage
    doc.add_heading("Application Usage Chart:", level=1)
    fig, ax = plt.subplots(figsize=(12, 10))  # Increased width and height for better visibility
    app_names = list(app_durations.keys())
    durations = list(app_durations.values())

    # Set up the bar chart
    y_pos = np.arange(len(app_names))
    ax.barh(y_pos, durations, color="skyblue")
    ax.set_yticks(y_pos)
    ax.set_yticklabels(app_names)
    ax.set_xlabel("Duration (seconds)")
    ax.set_title("Application Usage Time")

    # Rotate the y-axis labels
    plt.yticks(rotation=0)  # Keeping labels horizontal for better visibility

    # Add more space at the bottom
    plt.tight_layout(pad=4.0)

    # Save the chart as an image
    chart_path = "usage_chart.png"
    plt.savefig(chart_path)
    plt.close()

    # Add the chart image to the Word document
    doc.add_picture(chart_path, width=Inches(6))  # Adjust width if needed
    os.remove(chart_path)  # Clean up the chart image file

    # Save the Word document
    report_path = "Application_Usage_Report.docx"
    doc.save(report_path)
    print(f"Report generated: {report_path}")




# Step 4: Start and stop tracking functions
def start_tracking(icon, item):
    global start_time, tracking
    if not tracking:
        start_time = datetime.now()
        tracking = True
        print("Tracking started.")

def stop_tracking(icon, item):
    global start_time, tracking, work_duration
    if tracking and start_time:
        end_time = datetime.now()
        work_duration += end_time - start_time
        log_work_session(start_time, end_time)
        print(f"Tracking stopped. Session duration: {end_time - start_time}")
        start_time = None
        tracking = False

        # Fetch data and generate report after stopping
        logs = fetch_usage_data()  # Fetch data from the database
        if logs:
            generate_word_report(logs)
        else:
            print("No logs found to generate the report.")

# Step 5: Log work sessions in the database
def log_work_session(start, end):
    try:
        conn = pymysql.connect(**DB_CONFIG)
        with conn.cursor() as cursor:
            cursor.execute(
                "INSERT INTO work_logs (start_time, end_time) VALUES (%s, %s)",
                (start, end)
            )
            conn.commit()
    except pymysql.MySQLError as err:
        print(f"Error logging session: {err}")
    finally:
        if conn:
            conn.close()

# Step 6: Ensure the screenshot directory exists
def ensure_screenshot_directory():
    if not os.path.exists(screenshot_dir):
        os.makedirs(screenshot_dir)

# Step 7: Create a taskbar icon with options to start/stop tracking and exit
def create_tray_icon():
    try:
        # Load the icon image file
        icon_image = Image.open("E:/icon.jpg")  # Ensure 'icon.png' is in the same directory or provide the full path
    except FileNotFoundError:
        print("Icon file not found. Please ensure 'icon.png' is in the correct directory.")
        return None

    # Create the tray icon with menu items
    icon = pystray.Icon("Time Tracker", icon_image, title="Time Tracker")
    icon.menu = pystray.Menu(
        pystray.MenuItem("Start Tracking", start_tracking),
        pystray.MenuItem("Stop Tracking", stop_tracking),
        pystray.MenuItem("Exit", lambda icon, item: icon.stop())
    )
    return icon

# Step 8: Run screenshot capturing at intervals while tracking
def capture_screenshots():
    while True:
        if tracking:
            take_screenshot()
            time.sleep(interval)

# Step 9: Main function to initialize components and run the tracker
def main():
    ensure_screenshot_directory()
    icon = create_tray_icon()

    # Start capturing screenshots in a background thread
    screenshot_thread = threading.Thread(target=capture_screenshots)
    screenshot_thread.daemon = True
    screenshot_thread.start()

    icon.run()  # Display system tray icon

if __name__ == "__main__":
    main()
